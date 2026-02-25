#!/usr/bin/env python3
"""
Enhanced DOCX Converter - Razorpay Standard Format
Matches the formatting from: Tech Spec - Account Statements Read Cutoff

Key improvements:
- Arial font (matches Razorpay standard)
- Proper heading sizes: H1=18pt, H2=16pt, H3=14pt, H4=12pt
- Heading color: #1155cc (Razorpay blue)
- Body text: 11pt Arial black
- Title: 26pt
- Proper spacing matching Google Docs export
- Clean character encoding
- Better table formatting
"""

import os
import re
import subprocess
import sys
import tempfile
from pathlib import Path

# ── dependency check ──────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("python-docx not found. Install: pip install python-docx")
    sys.exit(1)

try:
    from bs4 import BeautifulSoup, NavigableString, Tag
except ImportError:
    print("beautifulsoup4 not found. Install: pip install beautifulsoup4")
    sys.exit(1)

# ── paths ─────────────────────────────────────────────────────────────────────
SCRIPT_DIR = Path(__file__).parent
INPUT_HTML = SCRIPT_DIR / "TDS_Payout_Integration_Tech_Spec_Summary.html"
OUTPUT_DOCX = SCRIPT_DIR / "TDS_Payout_Integration_Tech_Spec_Summary.docx"

# ── Razorpay Standard Colors (from benchmark doc) ────────────────────────────
RAZORPAY_BLUE = "1155cc"       # Headings color
TEXT_BLACK = "000000"           # Body text
LINK_BLUE = "0000ee"           # Hyperlinks
TABLE_BORDER = "CCCCCC"         # Table borders
INFO_BG = "EFF6FF"             # Info boxes
WARNING_BG = "FFFBEB"          # Warning boxes
SUCCESS_BG = "F0FDF4"          # Recommendation boxes
CODE_BG = "F8F9FA"             # Code blocks


# ─────────────────────────────────────────────────────────────────────────────
# 1. RENDER MERMAID DIAGRAMS
# ─────────────────────────────────────────────────────────────────────────────

def render_mermaid_diagrams(html: str, tmp_dir: str) -> tuple[str, list[str]]:
    """Find all mermaid divs, render to PNG, replace with img tags."""
    soup = BeautifulSoup(html, "html.parser")
    mermaid_divs = soup.find_all("div", class_="mermaid")
    image_paths = []

    for idx, div in enumerate(mermaid_divs):
        code = div.get_text().strip()
        mmd_path = os.path.join(tmp_dir, f"diagram_{idx}.mmd")
        png_path = os.path.join(tmp_dir, f"diagram_{idx}.png")

        with open(mmd_path, "w") as f:
            f.write(code)

        print(f"  Rendering diagram {idx + 1}/{len(mermaid_divs)}…")
        result = subprocess.run(
            ["npx", "-y", "@mermaid-js/mermaid-cli",
             "-i", mmd_path, "-o", png_path,
             "--backgroundColor", "white",
             "--width", "900"],
            capture_output=True, text=True, timeout=120
        )
        if result.returncode != 0 or not os.path.exists(png_path):
            print(f"    Warning: diagram {idx} failed to render.")
            new_tag = soup.new_tag("p")
            new_tag.string = f"[Diagram {idx + 1}: {code[:80]}…]"
            div.replace_with(new_tag)
        else:
            image_paths.append(png_path)
            img_tag = soup.new_tag("img",
                                   src=png_path,
                                   attrs={"class": "mermaid-rendered",
                                          "data-idx": str(idx)})
            div.replace_with(img_tag)

    return str(soup), image_paths


# ─────────────────────────────────────────────────────────────────────────────
# 2. DOCX FORMATTING HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_shading(cell, hex_color: str):
    """Apply background color to table cell."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def add_table_borders(table):
    """Add borders to all table cells."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), TABLE_BORDER)
        borders.append(el)
    tblPr.append(borders)


def add_paragraph_border(para, hex_color: str):
    """Add left border and shading for callout boxes."""
    pPr = para._p.get_or_add_pPr()
    # Shading
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)
    # Left border
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), RAZORPAY_BLUE)
    borders.append(left)
    pPr.append(borders)
    # Left indent
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "120")
    pPr.append(ind)


def set_run_font(run, size_pt=None, bold=False, italic=False,
                 color=None, mono=False):
    """Set font properties for a run."""
    if size_pt:
        run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if mono:
        run.font.name = "Courier New"
    else:
        run.font.name = "Arial"  # Razorpay standard


def add_hyperlink(paragraph, url, text, bold=False, italic=False, size_pt=11):
    """Add clickable hyperlink using Word's relationship mechanism."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Font
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')
    rPr.append(rFonts)

    # Blue color
    color = OxmlElement('w:color')
    color.set(qn('w:val'), LINK_BLUE)
    rPr.append(color)

    # Underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # Font size
    if size_pt:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(size_pt * 2))
        rPr.append(sz)
        sz_cs = OxmlElement('w:szCs')
        sz_cs.set(qn('w:val'), str(size_pt * 2))
        rPr.append(sz_cs)

    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    if italic:
        i = OxmlElement('w:i')
        rPr.append(i)

    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_horizontal_rule(doc):
    """Add a subtle horizontal line as section separator."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(12)

    # Add bottom border to create horizontal line
    pPr = para._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")  # 0.75pt line
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")  # Light gray
    borders.append(bottom)
    pPr.append(borders)


def add_toc_field(doc):
    """Insert Word TOC field."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)

    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    para._p.append(fld)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    para._p.append(instr)

    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    para._p.append(fld2)

    note = doc.add_paragraph()
    run = note.add_run("Right-click the TOC above → Update Field to refresh.")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string("94A3B8")


# ─────────────────────────────────────────────────────────────────────────────
# 3. INLINE TEXT PROCESSING
# ─────────────────────────────────────────────────────────────────────────────

def add_inline(para, element, base_size=11, bold=False, italic=False, mono=False):
    """Process inline elements (strong, em, code, a) recursively."""
    if isinstance(element, NavigableString):
        text = str(element)
        if text.strip():
            if mono:
                run = para.add_run(text)
                set_run_font(run, size_pt=base_size, bold=bold, italic=italic,
                            color=TEXT_BLACK, mono=True)
            else:
                run = para.add_run(text)
                set_run_font(run, size_pt=base_size, bold=bold, italic=italic,
                            color=TEXT_BLACK)
        return

    if isinstance(element, Tag):
        tag = element.name
        if tag == "strong" or tag == "b":
            for child in element.children:
                add_inline(para, child, base_size, True, italic, mono)
        elif tag == "em" or tag == "i":
            for child in element.children:
                add_inline(para, child, base_size, bold, True, mono)
        elif tag == "code":
            for child in element.children:
                add_inline(para, child, base_size, bold, italic, True)
        elif tag == "a":
            url = element.get("href", "")
            text = element.get_text()
            if url:
                add_hyperlink(para, url, text, bold=bold, italic=italic, size_pt=base_size)
            else:
                run = para.add_run(text)
                set_run_font(run, size_pt=base_size, bold=bold, italic=italic,
                            color=LINK_BLUE)
                run.font.underline = True
        else:
            for child in element.children:
                add_inline(para, child, base_size, bold, italic, mono)


# ─────────────────────────────────────────────────────────────────────────────
# 4. BLOCK ELEMENT PROCESSING
# ─────────────────────────────────────────────────────────────────────────────

class HeadingNumberer:
    """No-op: HTML headings already numbered."""
    def next(self, level: int) -> str:
        return ""


def process_element(element, doc, numberer):
    """Process HTML elements and convert to DOCX."""
    if isinstance(element, NavigableString):
        text = str(element).strip()
        if text:
            para = doc.add_paragraph()
            run = para.add_run(text)
            set_run_font(run, size_pt=11, color=TEXT_BLACK)
        return

    if not isinstance(element, Tag):
        return

    tag = element.name

    # ── Headings ──────────────────────────────────────────────────────────────
    if tag in ("h1", "h2", "h3", "h4"):
        level_map = {"h1": 1, "h2": 2, "h3": 3, "h4": 4}
        level = level_map[tag]
        para = doc.add_heading(level=level)
        para.paragraph_format.keep_with_next = True

        for child in element.children:
            add_inline(para, child, base_size=11)

        # Apply Razorpay standard heading formatting
        if para.runs:
            for run in para.runs:
                run.font.name = "Arial"
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(RAZORPAY_BLUE)

                # Size based on level (matching benchmark)
                if level == 1:
                    run.font.size = Pt(18)
                    para.paragraph_format.space_before = Pt(20)
                    para.paragraph_format.space_after = Pt(10)
                elif level == 2:
                    run.font.size = Pt(16)
                    para.paragraph_format.space_before = Pt(18)
                    para.paragraph_format.space_after = Pt(8)
                elif level == 3:
                    run.font.size = Pt(14)
                    para.paragraph_format.space_before = Pt(16)
                    para.paragraph_format.space_after = Pt(8)
                elif level == 4:
                    run.font.size = Pt(12)
                    para.paragraph_format.space_before = Pt(14)
                    para.paragraph_format.space_after = Pt(6)
        return

    # ── Paragraphs ────────────────────────────────────────────────────────────
    if tag == "p":
        classes = element.get("class", [])
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.line_spacing = 1.15

        for child in element.children:
            add_inline(para, child)

        # Apply Arial to all runs
        for run in para.runs:
            if not run.font.name or run.font.name == "Calibri":
                run.font.name = "Arial"
        return

    # ── Callout boxes ─────────────────────────────────────────────────────────
    if tag == "div":
        classes = element.get("class", [])

        if "recommendation" in classes:
            for child in element.children:
                if isinstance(child, Tag) and child.name in ("h3", "h4"):
                    para = doc.add_paragraph()
                    for c in child.children:
                        add_inline(para, c, base_size=12, bold=True)
                    add_paragraph_border(para, SUCCESS_BG)
                elif isinstance(child, Tag) and child.name == "p":
                    para = doc.add_paragraph()
                    for c in child.children:
                        add_inline(para, c)
                    add_paragraph_border(para, SUCCESS_BG)
            return

        if "warning" in classes:
            for child in element.children:
                if isinstance(child, Tag) and child.name in ("h3", "h4"):
                    para = doc.add_paragraph()
                    for c in child.children:
                        add_inline(para, c, base_size=12, bold=True)
                    add_paragraph_border(para, WARNING_BG)
                elif isinstance(child, Tag) and child.name == "p":
                    para = doc.add_paragraph()
                    for c in child.children:
                        add_inline(para, c)
                    add_paragraph_border(para, WARNING_BG)
            return

        if "info" in classes:
            for child in element.children:
                if isinstance(child, Tag) and child.name in ("h3", "h4"):
                    para = doc.add_paragraph()
                    for c in child.children:
                        add_inline(para, c, base_size=12, bold=True)
                    add_paragraph_border(para, INFO_BG)
                elif isinstance(child, Tag) and child.name == "p":
                    para = doc.add_paragraph()
                    for c in child.children:
                        add_inline(para, c)
                    add_paragraph_border(para, INFO_BG)
            return

        # Process children of other divs
        for child in element.children:
            process_element(child, doc, numberer)
        return

    # ── Lists ─────────────────────────────────────────────────────────────────
    if tag in ("ul", "ol"):
        for li in element.find_all("li", recursive=False):
            para = doc.add_paragraph(style="List Bullet" if tag == "ul" else "List Number")
            for child in li.children:
                if isinstance(child, Tag) and child.name in ("ul", "ol"):
                    continue
                add_inline(para, child)
            for run in para.runs:
                run.font.name = "Arial"
                run.font.size = Pt(11)
        return

    # ── Code blocks ───────────────────────────────────────────────────────────
    if tag == "pre":
        code_tag = element.find("code")
        code_text = code_tag.get_text() if code_tag else element.get_text()

        para = doc.add_paragraph()
        run = para.add_run(code_text)
        run.font.name = "Courier New"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(TEXT_BLACK)

        # Add background shading
        add_paragraph_border(para, CODE_BG)
        return

    # ── Tables ────────────────────────────────────────────────────────────────
    if tag == "table":
        rows = element.find_all("tr")
        if not rows:
            return

        cols = max(len(row.find_all(["th", "td"])) for row in rows)
        table = doc.add_table(rows=len(rows), cols=cols)
        add_table_borders(table)

        for row_idx, tr in enumerate(rows):
            cells = tr.find_all(["th", "td"])
            for col_idx, cell in enumerate(cells):
                doc_cell = table.rows[row_idx].cells[col_idx]

                # Header row
                if cell.name == "th":
                    set_cell_shading(doc_cell, RAZORPAY_BLUE)
                    para = doc_cell.paragraphs[0]
                    for child in cell.children:
                        add_inline(para, child, base_size=11, bold=True)
                    for run in para.runs:
                        run.font.color.rgb = RGBColor.from_string("FFFFFF")
                        run.font.name = "Arial"
                else:
                    para = doc_cell.paragraphs[0]
                    for child in cell.children:
                        add_inline(para, child)
                    for run in para.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(10)
        return

    # ── Images ────────────────────────────────────────────────────────────────
    if tag == "img":
        src = element.get("src")
        if src and os.path.exists(src):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(src, width=Inches(6))
        return

    # ── Horizontal rules ──────────────────────────────────────────────────────
    if tag == "hr":
        add_horizontal_rule(doc)
        return

    # ── Recursively process children ─────────────────────────────────────────
    for child in element.children:
        process_element(child, doc, numberer)


# ─────────────────────────────────────────────────────────────────────────────
# 5. MAIN CONVERSION
# ─────────────────────────────────────────────────────────────────────────────

def convert_html_to_docx(html_path: Path, output_path: Path):
    """Main conversion function."""
    with open(html_path, "r", encoding="utf-8") as f:
        html = f.read()

    print(f"\n{'='*60}")
    print("TDS Payout Summary → DOCX (Razorpay Standard Format)")
    print(f"{'='*60}\n")

    # Render Mermaid diagrams
    with tempfile.TemporaryDirectory() as tmp_dir:
        print(f"Step 1: Rendering Mermaid diagrams (tmp: {tmp_dir})…")
        html, image_paths = render_mermaid_diagrams(html, tmp_dir)

        print("\nStep 2: Building DOCX…")
        soup = BeautifulSoup(html, "html.parser")
        doc = Document()

        # ── Set document styles (Razorpay standard) ──────────────────────────
        # Normal style
        normal = doc.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(11)
        normal.font.color.rgb = RGBColor.from_string(TEXT_BLACK)
        normal.paragraph_format.space_after = Pt(8)
        normal.paragraph_format.line_spacing = 1.15

        # Heading styles (matching benchmark doc exactly)
        heading_config = {
            "Heading 1": 18,
            "Heading 2": 16,
            "Heading 3": 14,
            "Heading 4": 12,
        }

        for style_name, size in heading_config.items():
            try:
                h = doc.styles[style_name]
                h.font.name = "Arial"
                h.font.size = Pt(size)
                h.font.bold = True
                h.font.color.rgb = RGBColor.from_string(RAZORPAY_BLUE)
                h.paragraph_format.space_before = Pt(20 - (size - 12) // 2)
                h.paragraph_format.space_after = Pt(10 if size >= 16 else 8)
                h.paragraph_format.line_spacing = 1.15
            except KeyError:
                pass

        # ── Title block ───────────────────────────────────────────────────────
        title_para = doc.add_heading("TDS Payout Integration for BB+ Platform", 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.runs[0].font.name = "Arial"
        title_para.runs[0].font.size = Pt(26)  # Razorpay standard title size
        title_para.runs[0].font.color.rgb = RGBColor.from_string(TEXT_BLACK)
        title_para.runs[0].font.bold = True

        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = meta.add_run("Technical Specification — Summary Document")
        run.font.name = "Arial"
        run.font.size = Pt(13)
        run.font.italic = True
        run.font.color.rgb = RGBColor.from_string("475569")

        meta2 = doc.add_paragraph()
        meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = meta2.add_run("Author: Prashant Chauhan  ·  Version: 9.0  ·  February 2026")
        run2.font.name = "Arial"
        run2.font.size = Pt(11)
        run2.font.color.rgb = RGBColor.from_string("64748B")

        doc.add_paragraph()

        # ── TOC ───────────────────────────────────────────────────────────────
        toc_heading = doc.add_heading("Table of Contents", level=1)
        toc_heading.runs[0].font.name = "Arial"
        toc_heading.runs[0].font.size = Pt(14)
        toc_heading.runs[0].font.color.rgb = RGBColor.from_string(RAZORPAY_BLUE)
        add_toc_field(doc)
        doc.add_page_break()

        # ── Process body ──────────────────────────────────────────────────────
        body = soup.find("body") or soup
        numberer = HeadingNumberer()

        for child in body.children:
            process_element(child, doc, numberer)

        doc.save(str(output_path))
        print(f"\n✓ Saved: {output_path}")
        print(f"  Format: Razorpay Standard (Arial, H1=18pt, H2=16pt, H3=14pt)")
        print(f"  File size: {output_path.stat().st_size / 1024:.0f}KB")
        print(f"\nOpen: {output_path}")


if __name__ == "__main__":
    if not INPUT_HTML.exists():
        print(f"Error: {INPUT_HTML} not found")
        sys.exit(1)

    convert_html_to_docx(INPUT_HTML, OUTPUT_DOCX)
