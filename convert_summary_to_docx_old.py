#!/usr/bin/env python3
"""
Convert TDS Payout Integration Summary HTML → DOCX

Steps:
  1. Extract Mermaid diagrams → render to PNG via npx mmdc
  2. Replace mermaid divs with <img> placeholders in HTML
  3. Parse HTML sequentially → build DOCX with:
       - Title, Author, Date block
       - Word TOC field
       - Numbered headings (1. / 1.1. / 1.1.1.)
       - Real Word tables (blue header, bordered)
       - Callout boxes (info / warning / recommendation)
       - Code blocks (Courier New, shaded)
       - Embedded PNG images (Mermaid diagrams)
       - Bullet + numbered lists

Usage:
  python3 convert_summary_to_docx.py
"""

import os
import re
import subprocess
import sys
import tempfile
from pathlib import Path

# ── dependency check ──────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("python-docx not found. Install: pip install python-docx")
    sys.exit(1)

try:
    from bs4 import BeautifulSoup, NavigableString, Tag
except ImportError:
    print("beautifulsoup4 not found. Install: pip install beautifulsoup4")
    sys.exit(1)

# ── paths ─────────────────────────────────────────────────────────────────────
SCRIPT_DIR = Path(__file__).parent
INPUT_HTML = SCRIPT_DIR / "TDS_Payout_Integration_Tech_Spec_Summary.html"
OUTPUT_DOCX = SCRIPT_DIR / "TDS_Payout_Integration_Tech_Spec_Summary.docx"

# ── colours (Razorpay-ish palette) ───────────────────────────────────────────
BLUE_HEADER = "1A73E8"
BLUE_LIGHT  = "EFF6FF"
GREEN_LIGHT = "F0FDF4"
AMBER_LIGHT = "FFFBEB"
GRAY_CODE   = "F8F9FA"
DARK_TEXT   = "0F172A"
WHITE       = "FFFFFF"
BORDER_CLR  = "CCCCCC"

# ─────────────────────────────────────────────────────────────────────────────
# 1. RENDER MERMAID DIAGRAMS
# ─────────────────────────────────────────────────────────────────────────────

def render_mermaid_diagrams(html: str, tmp_dir: str) -> tuple[str, list[str]]:
    """
    Find all <div class="mermaid">…</div> blocks, render to PNG,
    replace with <img src="path"> and return modified HTML + image paths.
    """
    soup = BeautifulSoup(html, "html.parser")
    mermaid_divs = soup.find_all("div", class_="mermaid")
    image_paths = []

    for idx, div in enumerate(mermaid_divs):
        code = div.get_text().strip()
        mmd_path = os.path.join(tmp_dir, f"diagram_{idx}.mmd")
        png_path = os.path.join(tmp_dir, f"diagram_{idx}.png")

        with open(mmd_path, "w") as f:
            f.write(code)

        print(f"  Rendering diagram {idx + 1}/{len(mermaid_divs)}…")
        result = subprocess.run(
            ["npx", "-y", "@mermaid-js/mermaid-cli",
             "-i", mmd_path, "-o", png_path,
             "--backgroundColor", "white",
             "--width", "900"],
            capture_output=True, text=True, timeout=120
        )
        if result.returncode != 0 or not os.path.exists(png_path):
            print(f"    Warning: diagram {idx} failed to render. Inserting placeholder text.")
            new_tag = soup.new_tag("p")
            new_tag.string = f"[Diagram {idx + 1}: {code[:80]}…]"
            div.replace_with(new_tag)
        else:
            image_paths.append(png_path)
            img_tag = soup.new_tag("img",
                                   src=png_path,
                                   attrs={"class": "mermaid-rendered",
                                          "data-idx": str(idx)})
            div.replace_with(img_tag)

    return str(soup), image_paths


# ─────────────────────────────────────────────────────────────────────────────
# 2. DOCX HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_shading(cell, hex_color: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def add_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), BORDER_CLR)
        borders.append(el)
    tblPr.append(borders)


def add_paragraph_border(para, hex_color: str):
    """Left border + light shaded background for callout boxes."""
    pPr = para._p.get_or_add_pPr()
    # shading
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)
    # left accent border
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), BLUE_HEADER)
    borders.append(left)
    pPr.append(borders)
    # modest left indent to clear the border (120 twips ≈ 0.08 inch)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "120")
    pPr.append(ind)


def set_run_font(run, size_pt=None, bold=False, italic=False,
                 color=None, mono=False):
    if size_pt:
        run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if mono:
        run.font.name = "Courier New"


def add_toc_field(doc):
    """Insert a Word TOC field (updates on Ctrl+A, F9 in Word)."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    para._p.append(fld)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    para._p.append(instr)

    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    para._p.append(fld2)

    note = doc.add_paragraph()
    run = note.add_run("Right-click the TOC above → Update Field to refresh page numbers.")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string("94A3B8")


# ─────────────────────────────────────────────────────────────────────────────
# 3. HEADING NUMBERING (kept as stub — HTML headings are already numbered)
# ─────────────────────────────────────────────────────────────────────────────

class HeadingNumberer:
    """No-op: HTML headings already contain their numbers."""
    def next(self, level: int) -> str:
        return ""


# ─────────────────────────────────────────────────────────────────────────────
# 4. INLINE TEXT HELPER
# ─────────────────────────────────────────────────────────────────────────────

def add_hyperlink(paragraph, url, text, bold=False, italic=False, size_pt=11):
    """
    Add a clickable hyperlink to a paragraph.
    Uses Word's hyperlink relationship mechanism.
    """
    # Get the document part
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the w:hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a new run for the hyperlink text
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Hyperlink styling - blue and underlined
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '1A73E8')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # Font size
    if size_pt:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(size_pt * 2))  # Word uses half-points
        rPr.append(sz)
        sz_cs = OxmlElement('w:szCs')
        sz_cs.set(qn('w:val'), str(size_pt * 2))
        rPr.append(sz_cs)

    # Bold/italic if needed
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    if italic:
        i = OxmlElement('w:i')
        rPr.append(i)

    new_run.append(rPr)

    # Add the text
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_inline(para, element, base_bold=False, base_italic=False,
               base_size=11, base_mono=False):
    """Recursively add inline content (strong, em, code, a, text) to para."""
    if isinstance(element, NavigableString):
        # Collapse newlines/indents inside HTML to a single space
        text = re.sub(r'\s+', ' ', str(element))
        if text.strip():
            run = para.add_run(text)
            set_run_font(run, size_pt=base_size, bold=base_bold,
                         italic=base_italic, mono=base_mono)
        return

    if not isinstance(element, Tag):
        return

    tag = element.name
    bold   = base_bold   or tag in ("strong", "b")
    italic = base_italic or tag in ("em", "i")
    mono   = base_mono   or tag == "code"

    if tag == "br":
        para.add_run("\n")
        return
    if tag == "a":
        url = element.get("href", "")
        text = element.get_text()
        if url:
            # Create actual clickable hyperlink
            add_hyperlink(para, url, text, bold=bold, italic=italic, size_pt=base_size)
        else:
            # Fallback to styled text if no href
            run = para.add_run(text)
            set_run_font(run, size_pt=base_size, bold=bold, italic=italic,
                         color="1A73E8")
            run.font.underline = True
        return

    for child in element.children:
        add_inline(para, child, bold, italic, base_size, mono)


# ─────────────────────────────────────────────────────────────────────────────
# 5. MAIN ELEMENT PROCESSOR
# ─────────────────────────────────────────────────────────────────────────────

def process_element(el, doc, numberer: HeadingNumberer):
    if isinstance(el, NavigableString):
        return
    if not isinstance(el, Tag):
        return

    tag    = el.name
    classes = set(el.get("class", []))

    # ── Headings ──────────────────────────────────────────────────────────────
    heading_map = {"h1": 1, "h2": 2, "h3": 3, "h4": 4}
    if tag in heading_map:
        level = heading_map[tag]
        text  = el.get_text(strip=True)
        # Skip the big HTML doc-title — already written as Word Title style
        if level == 1:
            return
        # h4 → bold paragraph (12pt dark), not a Word Heading style
        if level == 4:
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(8)
            para.paragraph_format.space_after  = Pt(2)
            run = para.add_run(text)
            set_run_font(run, size_pt=12, bold=True, color=DARK_TEXT)
            return
        # HTML headings already contain numbers (e.g. "1.3 The Problem")
        # — use text as-is, no additional auto-numbering
        doc.add_heading(text, level=level)
        return

    # ── Horizontal rule ───────────────────────────────────────────────────────
    if tag == "hr":
        para = doc.add_paragraph()
        pPr  = para._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "6")
        bot.set(qn("w:color"), "E2E8F0")
        borders.append(bot)
        pPr.append(borders)
        return

    # ── Callout boxes ─────────────────────────────────────────────────────────
    if "info" in classes or "recommendation" in classes or "warning" in classes:
        color = (AMBER_LIGHT if "warning" in classes else
                 GREEN_LIGHT if "recommendation" in classes else
                 BLUE_LIGHT)
        # heading inside callout box
        title_el = el.find(["h4", "h3", "strong"])
        if title_el:
            para = doc.add_paragraph()
            add_paragraph_border(para, color)
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after  = Pt(2)
            run = para.add_run(title_el.get_text(strip=True))
            set_run_font(run, size_pt=12, bold=True, color=DARK_TEXT)

        def _callout_para(tight=False):
            p = doc.add_paragraph()
            add_paragraph_border(p, color)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(2 if not tight else 0)
            return p

        # body text (excluding the title element)
        for child in el.children:
            if isinstance(child, NavigableString):
                txt = str(child).strip()
                if txt:
                    p = _callout_para()
                    run = p.add_run(txt)
                    run.font.size = Pt(11)
                continue
            if not isinstance(child, Tag):
                continue
            if child == title_el:
                continue
            if child.name == "p":
                p = _callout_para()
                for c in child.children:
                    add_inline(p, c, base_size=11)
            elif child.name in ("ul", "ol"):
                for li in child.find_all("li", recursive=False):
                    # manual bullet to avoid List Bullet style indentation clash
                    p = _callout_para()
                    pPr = p._p.get_or_add_pPr()
                    # hanging indent: first line starts at left, wrap indents 180 twips
                    ind = p._p.get_or_add_pPr().find(qn("w:ind"))
                    if ind is None:
                        ind = OxmlElement("w:ind")
                        pPr.append(ind)
                    ind.set(qn("w:left"),    "300")
                    ind.set(qn("w:hanging"), "180")
                    bullet_run = p.add_run("•  ")
                    bullet_run.font.size = Pt(11)
                    for c in li.children:
                        if isinstance(c, Tag) and c.name in ("ul", "ol"):
                            continue  # skip nested for now
                        add_inline(p, c, base_size=11)

        # trailing spacer after callout block
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after  = Pt(4)
        return

    # ── Images (rendered Mermaid) ─────────────────────────────────────────────
    if tag == "img":
        src = el.get("src", "")
        if src and os.path.exists(src):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            try:
                run.add_picture(src, width=Inches(6))
            except Exception as e:
                run.text = f"[Image: {os.path.basename(src)}]"
            doc.add_paragraph()  # spacing
        return

    # ── mermaid container (fallback if not pre-rendered) ─────────────────────
    if "mermaid-container" in classes:
        for child in el.children:
            process_element(child, doc, numberer)
        return

    # ── Code blocks ───────────────────────────────────────────────────────────
    if tag == "pre":
        code_el = el.find("code")
        text = (code_el or el).get_text()
        for line in text.split("\n"):
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after  = Pt(0)
            pPr = para._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), GRAY_CODE)
            pPr.append(shd)
            # 180 twips = 0.125 inch left indent — just enough to look nested
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), "180")
            pPr.append(ind)
            run = para.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        doc.add_paragraph()  # breathing room after code block
        return

    # ── Tables ────────────────────────────────────────────────────────────────
    if tag == "table":
        rows = el.find_all("tr")
        if not rows:
            return
        max_cols = max(len(r.find_all(["td", "th"])) for r in rows)
        if max_cols == 0:
            return
        table = doc.add_table(rows=0, cols=max_cols)
        table.style = "Table Grid"
        add_table_borders(table)

        for r_idx, tr in enumerate(rows):
            cells = tr.find_all(["th", "td"])
            row   = table.add_row()
            for c_idx, cell_el in enumerate(cells):
                if c_idx >= max_cols:
                    break
                cell = row.cells[c_idx]
                is_header = (cell_el.name == "th" or r_idx == 0)
                if is_header:
                    set_cell_shading(cell, BLUE_HEADER)
                else:
                    set_cell_shading(cell, "F8FAFC" if r_idx % 2 == 0 else WHITE)

                # nested list inside cell
                lists = cell_el.find_all(["ul", "ol"])
                if lists:
                    cell.paragraphs[0].clear()
                    for lst in lists:
                        for li in lst.find_all("li", recursive=False):
                            p = cell.add_paragraph(style="List Bullet")
                            for c in li.children:
                                add_inline(p, c, base_size=11,
                                           base_bold=is_header)
                else:
                    para = cell.paragraphs[0]
                    for c in cell_el.children:
                        add_inline(para, c, base_size=11, base_bold=is_header)
                    if is_header:
                        for run in para.runs:
                            run.font.color.rgb = RGBColor.from_string(WHITE)
        doc.add_paragraph()
        return

    # ── Lists ─────────────────────────────────────────────────────────────────
    if tag in ("ul", "ol"):
        style = "List Bullet" if tag == "ul" else "List Number"
        for li in el.find_all("li", recursive=False):
            para = doc.add_paragraph(style=style)
            for c in li.children:
                if isinstance(c, Tag) and c.name in ("ul", "ol"):
                    process_element(c, doc, numberer)
                else:
                    add_inline(para, c, base_size=11)
        return

    # ── Paragraphs ────────────────────────────────────────────────────────────
    if tag == "p":
        text = el.get_text(strip=True)
        if not text:
            return
        para = doc.add_paragraph()
        for c in el.children:
            add_inline(para, c, base_size=11)
        return

    # ── approach-header badges (skip decoration, just add text) ───────────────
    if "approach-header" in classes:
        badge = el.find(class_=re.compile("badge"))
        muted = el.find(class_="muted")
        if badge or muted:
            parts = []
            if badge: parts.append(badge.get_text(strip=True))
            if muted: parts.append(muted.get_text(strip=True))
            para = doc.add_paragraph(" — ".join(parts))
            para.runs[0].font.size = Pt(10)
            para.runs[0].font.italic = True
        return

    # ── skip nav/script/style elements ───────────────────────────────────────
    if tag in ("script", "style", "nav", "head"):
        return

    # ── recurse into container divs/sections ─────────────────────────────────
    skip_classes = {"toc", "toc-box"}
    if classes & skip_classes:
        return

    for child in el.children:
        process_element(child, doc, numberer)


# ─────────────────────────────────────────────────────────────────────────────
# 6. BUILD DOCX
# ─────────────────────────────────────────────────────────────────────────────

def build_docx(html: str, output_path: Path):
    soup = BeautifulSoup(html, "html.parser")
    doc  = Document()

    # Page margins (2.5 cm each side)
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Default Normal style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)

    # Heading styles — explicit sizes matching Razorpay tech spec template
    heading_cfg = {
        "Heading 1": (16, "1A3C6E"),
        "Heading 2": (14, "1A3C6E"),
        "Heading 3": (13, "334155"),
    }
    for style_name, (size, color_hex) in heading_cfg.items():
        try:
            h = doc.styles[style_name]
            h.font.name = "Calibri"
            h.font.size = Pt(size)
            h.font.bold = True
            h.font.color.rgb = RGBColor.from_string(color_hex)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after  = Pt(4)
        except KeyError:
            pass

    # ── Title block ───────────────────────────────────────────────────────────
    title_para = doc.add_heading("TDS Payout Integration for BB+ Platform", 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.runs[0].font.size = Pt(24)
    title_para.runs[0].font.color.rgb = RGBColor.from_string("1A3C6E")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Technical Specification — Summary Document")
    run.font.size = Pt(13)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string("475569")

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = meta2.add_run("Author: Prashant Chauhan  ·  Version: 9.0  ·  February 2026")
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor.from_string("64748B")

    doc.add_paragraph()

    # ── TOC ───────────────────────────────────────────────────────────────────
    toc_heading = doc.add_heading("Table of Contents", level=1)
    toc_heading.runs[0].font.size = Pt(14)
    add_toc_field(doc)
    doc.add_page_break()

    # ── Process body ─────────────────────────────────────────────────────────
    body = soup.find("body") or soup
    numberer = HeadingNumberer()

    for child in body.children:
        process_element(child, doc, numberer)

    doc.save(str(output_path))
    print(f"\n✓ Saved: {output_path}")


# ─────────────────────────────────────────────────────────────────────────────
# 7. ENTRY POINT
# ─────────────────────────────────────────────────────────────────────────────

def main():
    print("=== TDS Payout Summary → DOCX ===\n")

    html = INPUT_HTML.read_text(encoding="utf-8")

    with tempfile.TemporaryDirectory() as tmp:
        print(f"Step 1: Rendering Mermaid diagrams (tmp: {tmp})…")
        html_modified, images = render_mermaid_diagrams(html, tmp)

        print(f"\nStep 2: Building DOCX…")
        build_docx(html_modified, OUTPUT_DOCX)

    print(f"\nDone. Open: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
